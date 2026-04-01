from __future__ import annotations

import logging
from pathlib import Path

from langchain_community.document_loaders import (
    PyMuPDFLoader,
    UnstructuredMarkdownLoader,
)
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md"}


def _load_pdf_unstructured(file_path: str) -> list[Document]:
    """Primary PDF loader using unstructured for layout-aware extraction."""
    from unstructured.partition.pdf import partition_pdf

    elements = partition_pdf(str(file_path), strategy="hi_res")
    docs: list[Document] = []
    for el in elements:
        meta = {
            "element_type": type(el).__name__,
            "page": getattr(el.metadata, "page_number", None),
        }
        parent = getattr(el.metadata, "parent_id", None)
        if parent:
            meta["parent_element_id"] = parent
        docs.append(Document(page_content=str(el), metadata=meta))
    return docs


def _load_pdf_fallback(file_path: str) -> list[Document]:
    """Fallback PDF loader using PyMuPDF."""
    loader = PyMuPDFLoader(file_path)
    return loader.load()


def _load_pdf(file_path: str) -> list[Document]:
    try:
        return _load_pdf_unstructured(file_path)
    except Exception:
        logger.warning(
            "Unstructured PDF parsing failed for %s, falling back to PyMuPDF",
            file_path,
            exc_info=True,
        )
        return _load_pdf_fallback(file_path)


def _load_docx(file_path: str) -> list[Document]:
    """Load DOCX with heading-aware structure via python-docx."""
    from docx import Document as DocxDocument

    docx_doc = DocxDocument(file_path)
    docs: list[Document] = []
    current_section = ""
    current_parent_section = ""
    buffer: list[str] = []

    def _flush() -> None:
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={
                    "element_type": "NarrativeText",
                    "section_title": current_section,
                    "parent_section": current_parent_section,
                },
            ))
        buffer.clear()

    for para in docx_doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading 1") or style_name == "title":
            _flush()
            current_parent_section = ""
            current_section = para.text.strip()
        elif style_name.startswith("heading 2"):
            _flush()
            current_parent_section = current_section
            current_section = para.text.strip()
        elif style_name.startswith("heading"):
            _flush()
            current_section = para.text.strip()
        else:
            text = para.text.strip()
            if text:
                buffer.append(text)

    _flush()

    for table in docx_doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            docs.append(Document(
                page_content=table_text,
                metadata={
                    "element_type": "Table",
                    "section_title": current_section,
                    "parent_section": current_parent_section,
                },
            ))

    return docs


def _load_markdown(file_path: str) -> list[Document]:
    loader = UnstructuredMarkdownLoader(file_path)
    return loader.load()


_LOADER_DISPATCH = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".md": _load_markdown,
}


def load_document(
    file_path: str,
    doc_type: str | None = None,
    original_name: str | None = None,
) -> list[Document]:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in _LOADER_DISPATCH:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    loader_fn = _LOADER_DISPATCH[ext]
    docs = loader_fn(str(path))

    for doc in docs:
        doc.metadata["source_file"] = original_name or path.name
        doc.metadata["file_extension"] = ext
        if doc_type:
            doc.metadata["doc_type"] = doc_type

    return docs
